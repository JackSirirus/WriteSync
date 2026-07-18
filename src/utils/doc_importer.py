"""
Doc Importer — 文档导入与结构化解析模块

DocumentParser class with methods to parse .md, .txt, .docx files
and split into structured data (chapters, settings, materials).

Supports:
- Markdown: split by # / ## headers
- Plaintext: detect 第X章, Chapter X, 一、二、 numbered sections
- DOCX: python-docx with zipfile+xml fallback

Output: ImportResult dataclass
- chapters: list[{title, content, chapter_num}]
- settings: list[{type, name, content}]
- materials: list[{title, content, tags}]
- metadata: {source_filename, total_chars, detected_format, word_count, file_size}
"""

import logging
import re
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger("writesync")

# ═══════════════════════════════════════════════════════════════
# Data Models
# ═══════════════════════════════════════════════════════════════

# 章节分隔符正则（plaintext）
_CHAPTER_PATTERNS: list[re.Pattern] = [
    re.compile(r'第[一二三四五六七八九十百千\d]+[章节]'),  # 第X章 / 第X节
    re.compile(r'第\s*(\d+)\s*[章节]'),                 # 第 1 章
    re.compile(r'\bChapter\s+\d+', re.IGNORECASE),      # Chapter X
    re.compile(r'\bPart\s+\d+', re.IGNORECASE),         # Part X
    re.compile(r'^[一二三四五六七八九十]+[、，,]', re.MULTILINE),        # 一、二、三、
    re.compile(r'^[（(][一二三四五六七八九十]+[)）]', re.MULTILINE),      # (一) (二)
    re.compile(r'^\d+[\.\、]\s', re.MULTILINE),                       # 1. 2. 3.
]

# 设定关键词（用于归类为 settings）
_SETTING_KEYWORDS: list[str] = [
    "世界观", "设定", "世界设定", "世界背景",
    "力量体系", "能力体系", "修炼体系", "魔法体系",
    "能力", "技能", "功法", "异能",
    "种族", "势力", "阵营", "组织",
    "地理", "版图", "地图", "区域",
]

# 材料自动标签映射
_MATERIAL_TAG_HINTS: dict[str, list[str]] = {
    "角色": ["角色", "人物", "主角", "反派", "配角", "性格", "身份", "关系"],
    "世界观": ["世界", "地理", "设定", "背景", "体系", "规则", "环境"],
    "情节": ["情节", "剧情", "故事", "事件", "冲突", "转折", "高潮", "结局"],
    "参考": ["参考", "研究", "资料", "来源", "数据", "分析", "调研", "笔记"],
    "灵感": ["灵感", "想法", "创意", "点子", "脑洞"],
}


@dataclass
class ImportResult:
    """文档导入结果"""

    chapters: list[dict] = field(default_factory=list)   # {title, content, chapter_num}
    settings: list[dict] = field(default_factory=list)   # {type, name, content}
    materials: list[dict] = field(default_factory=list)  # {title, content, tags}
    metadata: dict = field(default_factory=dict)         # source_filename, total_chars, etc.

    def to_dict(self) -> dict:
        return {
            "chapters": self.chapters,
            "settings": self.settings,
            "materials": self.materials,
            "metadata": self.metadata,
        }

    @classmethod
    def from_dict(cls, d: dict) -> "ImportResult":
        return cls(
            chapters=d.get("chapters", []),
            settings=d.get("settings", []),
            materials=d.get("materials", []),
            metadata=d.get("metadata", {}),
        )

    def is_empty(self) -> bool:
        return not self.chapters and not self.settings and not self.materials


@dataclass
class _Section:
    """内部中间结构：一个解析出的段落/区块"""
    title: str
    content: str
    index: int = 0  # 原始索引


# ═══════════════════════════════════════════════════════════════
# DocumentParser
# ═══════════════════════════════════════════════════════════════

class DocumentParser:
    """文档解析器：将 .md / .txt / .docx 文件解析为 ImportResult。"""

    LARGE_FILE_THRESHOLD = 100 * 1024  # 100KB

    # ── Public API ──

    def parse(self, content: str, filename: str = "") -> ImportResult:
        """主入口：解析文档内容并返回 ImportResult。

        Args:
            content: 文档文本内容（或 docx 文件路径）
            filename: 源文件名（用于判断格式）

        Returns:
            ImportResult 结构化解析结果
        """
        file_ext = Path(filename).suffix.lower() if filename else ""

        # 检查是否可能是二进制 docx 路径
        if file_ext == ".docx":
            return self._parse_docx(content if Path(content).is_file() else filename)

        if file_ext in (".md", ".markdown"):
            result = self._parse_markdown(content)
        else:
            result = self._parse_plaintext(content)

        # 填入元数据
        result.metadata["source_filename"] = filename or "(raw text)"
        result.metadata["total_chars"] = len(content)
        result.metadata["detected_format"] = file_ext.lstrip(".") if file_ext else "text"

        # 字数统计
        result.metadata["word_count"] = _count_chars(content)

        # 大文件警告
        byte_size = len(content.encode("utf-8", errors="replace"))
        result.metadata["file_size"] = _readable_size_bytes(byte_size)
        if byte_size > self.LARGE_FILE_THRESHOLD:
            logger.warning(
                "DocumentParser: file '%s' is %.1fKB (>100KB), processing may be slow",
                filename, byte_size / 1024,
            )

        return result

    def parse_file(self, file_path: str) -> ImportResult:
        """从文件路径解析。

        Args:
            file_path: 文件路径

        Returns:
            ImportResult，若文件不存在/无法读取则返回含 error 的 result
        """
        fpath = Path(file_path)
        if not fpath.exists():
            logger.error("DocumentParser: file not found: %s", file_path)
            return ImportResult(
                metadata={"source_filename": fpath.name, "error": f"文件不存在: {file_path}"},
            )

        suffix = fpath.suffix.lower()

        # 二进制文件检查（非 docx）
        if suffix not in (".md", ".markdown", ".txt", ".docx"):
            logger.warning("DocumentParser: unsupported file type: %s", suffix)
            return ImportResult(
                metadata={
                    "source_filename": fpath.name,
                    "error": f"不支持的文件类型: {suffix}（支持 .md .txt .docx）",
                    "file_size": _readable_size_bytes(fpath.stat().st_size),
                },
            )

        # docx 走二进制解析
        if suffix == ".docx":
            return self._parse_docx(str(fpath))

        try:
            content = fpath.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            logger.error("DocumentParser: failed to read file: %s", e)
            return ImportResult(
                metadata={"source_filename": fpath.name, "error": f"读取文件失败: {e}"},
            )

        return self.parse(content, fpath.name)

    # ── Format Parsers ──

    def _parse_markdown(self, content: str) -> ImportResult:
        """解析 Markdown：按 # / ## / ### 标题分割。"""
        if not content or not content.strip():
            return ImportResult()

        sections = _split_by_markdown_headings(content)

        if not sections:
            # 无标题 → 整体作为材料
            return ImportResult(
                materials=[{
                    "title": "(无标题)",
                    "content": content,
                    "tags": _auto_tag(content, ""),
                }],
            )

        return _classify_sections(sections)

    def _parse_plaintext(self, content: str) -> ImportResult:
        """解析纯文本：按「第X章」/ 空行 / 编号段落分割。"""
        if not content or not content.strip():
            return ImportResult()

        sections = _split_by_chapter_patterns(content)

        if not sections:
            # 无章节结构 → 尝试按空行分块
            sections = _split_by_blank_lines(content)

        if not sections:
            return ImportResult(
                materials=[{
                    "title": "(无标题)",
                    "content": content,
                    "tags": _auto_tag(content, ""),
                }],
            )

        return _classify_sections(sections)

    def _parse_docx(self, file_path: str) -> ImportResult:
        """解析 .docx：优先 python-docx，fallback 到 zipfile+xml。"""
        fpath = Path(file_path)
        if not fpath.is_file():
            return ImportResult(
                metadata={
                    "source_filename": fpath.name,
                    "error": f"文件不存在: {file_path}",
                },
            )

        file_size = fpath.stat().st_size

        # 尝试 python-docx
        try:
            from docx import Document  # type: ignore[import-untyped]
            return self._parse_docx_via_library(file_path, fpath, file_size)
        except ImportError:
            logger.info("DocumentParser: python-docx not installed, using zipfile fallback")
        except Exception as e:
            logger.warning("DocumentParser: python-docx failed (%s), trying zipfile fallback", e)

        # zipfile + xml fallback
        return self._parse_docx_via_zipfile(file_path, fpath, file_size)

    def _parse_docx_via_library(
        self, file_path: str, fpath: Path, file_size: int,
    ) -> ImportResult:
        """通过 python-docx 解析。"""
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(str(fpath))
        sections_raw: list[_Section] = []
        current_title = ""
        current_body_lines: list[str] = []
        idx = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = (
                para.style
                and para.style.name
                and para.style.name.startswith("Heading")
            )

            if is_heading and current_body_lines:
                # 保存上一个 section
                sections_raw.append(_Section(
                    title=current_title or f"章节{idx+1}",
                    content="\n".join(current_body_lines),
                    index=idx,
                ))
                idx += 1
                current_title = text
                current_body_lines = []

            if is_heading:
                current_title = text
            else:
                current_body_lines.append(text)

        # 最后一个 section
        if current_body_lines:
            sections_raw.append(_Section(
                title=current_title or f"章节{idx+1}",
                content="\n".join(current_body_lines),
                index=idx,
            ))

        # 无标题段落时用全文字段
        if not sections_raw:
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if not full_text.strip():
                return ImportResult()
            return ImportResult(
                materials=[{
                    "title": fpath.stem,
                    "content": full_text,
                    "tags": _auto_tag(full_text, fpath.stem),
                }],
            )

        return _finalize_import(sections_raw, fpath, file_size, "docx")

    def _parse_docx_via_zipfile(
        self, file_path: str, fpath: Path, file_size: int,
    ) -> ImportResult:
        """通过 zipfile + xml.etree 解析 docx 的 word/document.xml。"""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                if "word/document.xml" not in z.namelist():
                    return ImportResult(
                        metadata={
                            "source_filename": fpath.name,
                            "error": "无效的 .docx 文件：缺少 word/document.xml",
                        },
                    )
                xml_bytes = z.read("word/document.xml")
        except zipfile.BadZipFile:
            return ImportResult(
                metadata={
                    "source_filename": fpath.name,
                    "error": "无效的 .docx 文件：无法作为 ZIP 打开",
                },
            )
        except Exception as e:
            return ImportResult(
                metadata={
                    "source_filename": fpath.name,
                    "error": f"读取 .docx 失败: {e}",
                },
            )

        # 解析 XML 提取文本
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        paragraphs: list[str] = []
        for p_elem in root.iter(f"{{{ns['w']}}}p"):
            texts: list[str] = []
            for t_elem in p_elem.iter(f"{{{ns['w']}}}t"):
                if t_elem.text:
                    texts.append(t_elem.text)
            line = "".join(texts).strip()
            if line:
                paragraphs.append(line)

        full_text = "\n".join(paragraphs)
        if not full_text.strip():
            return ImportResult()

        # 尝试按标题样式解析
        sections_raw: list[_Section] = []
        current_title = ""
        current_body_lines: list[str] = []

        for para in paragraphs:
            # 检查是否是加粗/大字号标题（启发式：短行、无标点结尾）
            is_likely_heading = (
                len(para) <= 50
                and not para.endswith(("。", "，", "；", "：", "、"))
                and not para.endswith((".", ",", ";", ":", "-"))
            )
            if is_likely_heading and current_body_lines:
                sections_raw.append(_Section(
                    title=current_title or "章节",
                    content="\n".join(current_body_lines),
                    index=len(sections_raw),
                ))
                current_title = para
                current_body_lines = []
                continue

            if is_likely_heading and not current_body_lines:
                current_title = para
                continue

            current_body_lines.append(para)

        if current_body_lines:
            sections_raw.append(_Section(
                title=current_title or "章节",
                content="\n".join(current_body_lines),
                index=len(sections_raw),
            ))

        if not sections_raw:
            return ImportResult(
                materials=[{
                    "title": fpath.stem,
                    "content": full_text,
                    "tags": _auto_tag(full_text, fpath.stem),
                }],
            )

        return _finalize_import(sections_raw, fpath, file_size, "docx")


# ═══════════════════════════════════════════════════════════════
# Splitting Helpers
# ═══════════════════════════════════════════════════════════════

def _split_by_markdown_headings(content: str) -> list[_Section]:
    """按 Markdown 标题（# / ## / ###）分割。"""
    heading_pat = re.compile(r'^(#{1,3})\s+(.*?)$', re.MULTILINE)
    headings: list[tuple[int, int, str]] = []  # (start_pos, level, title)

    for m in heading_pat.finditer(content):
        level = len(m.group(1))
        title = m.group(2).strip()
        headings.append((m.start(), level, title))

    if not headings:
        return []

    sections: list[_Section] = []
    for i, (pos, _level, title) in enumerate(headings):
        end = headings[i + 1][0] if i + 1 < len(headings) else len(content)
        body = content[pos:end]

        # 去掉标题行本身
        header_line = re.match(r'^#{1,3}\s+.*$', body, re.MULTILINE)
        if header_line:
            body = body[header_line.end():].strip()

        sections.append(_Section(title=title, content=body, index=i))

    return sections


def _split_by_chapter_patterns(content: str) -> list[_Section]:
    """按中文章节分隔符（第X章 / 一、二、/ Chapter X 等）分割。"""
    # 合并所有模式找到所有匹配位置
    all_matches: list[tuple[int, int, str]] = []  # (start, end, matched_text)

    for pattern in _CHAPTER_PATTERNS:
        for m in pattern.finditer(content):
            all_matches.append((m.start(), m.end(), m.group().strip()))

    if not all_matches:
        return []

    # 去重并按位置排序
    all_matches.sort()
    unique: list[tuple[int, int, str]] = []
    seen_ranges: set[tuple[int, int]] = set()
    for start, end, text in all_matches:
        # 合并重叠匹配（同一行）
        normalized = (start // 80 * 80, end)  # 粗略去重：80 字符窗口
        if normalized not in seen_ranges:
            seen_ranges.add(normalized)
            unique.append((start, end, text))

    unique.sort()

    sections: list[_Section] = []
    for i, (start, end, match_text) in enumerate(unique):
        # 取从匹配位置到下一匹配位置（或结尾）的内容
        next_start = unique[i + 1][0] if i + 1 < len(unique) else len(content)
        section_body = content[start:next_start].strip()

        # 第一行是标题（匹配到的分隔符行），后面是正文
        first_newline = section_body.find("\n")
        if first_newline > 0:
            title = section_body[:first_newline].strip()
            body = section_body[first_newline + 1:].strip()
        else:
            title = match_text
            body = ""

        sections.append(_Section(title=title, content=body, index=i))

    return sections


def _split_by_blank_lines(content: str) -> list[_Section]:
    """按空行分块，取首行作标题。"""
    blocks = [b.strip() for b in re.split(r'\n\s*\n', content) if b.strip()]
    if len(blocks) <= 1:
        return []

    sections: list[_Section] = []
    for i, block in enumerate(blocks):
        lines = block.split("\n")
        title = lines[0].strip()[:50]  # 首行作标题（限 50 字）
        body = block
        sections.append(_Section(title=title, content=body, index=i))

    return sections


# ═══════════════════════════════════════════════════════════════
# Classification Helpers
# ═══════════════════════════════════════════════════════════════

def _classify_sections(sections: list[_Section]) -> ImportResult:
    """将中间 section 列表分类为 chapters / settings / materials。"""
    chapters: list[dict] = []
    settings: list[dict] = []
    materials: list[dict] = []

    for sec in sections:
        combined_text = sec.title + " " + sec.content[:200]

        # Chapter detection HAS PRIORITY over settings detection.
        # A heading like "第一章 觉醒" with "能力" in content is a chapter, not a setting.
        if _looks_like_chapter(sec.title, combined_text):
            chapters.append({
                "title": sec.title,
                "content": sec.content,
                "chapter_num": len(chapters) + 1,
            })
        elif _is_setting(combined_text):
            settings.append({
                "type": "worldbuilding",
                "name": sec.title,
                "content": sec.content,
            })
        else:
            materials.append({
                "title": sec.title,
                "content": sec.content,
                "tags": _auto_tag(sec.content, sec.title),
            })

    # 如果全部是 chapters 但没有内容被识别为 setting/material → 无材料
    # 如果 chapters 为空但 sections 非空 → 全部归为 materials
    if not chapters and not settings and materials:
        # 尝试把 materials 中有章节特征的提升为 chapters
        new_chapters: list[dict] = []
        remaining: list[dict] = []
        for m in materials:
            if _looks_like_chapter(m["title"], m["content"][:200]):
                new_chapters.append({
                    "title": m["title"],
                    "content": m["content"],
                    "chapter_num": len(new_chapters) + 1,
                })
            else:
                remaining.append(m)
        if new_chapters:
            chapters = new_chapters
            materials = remaining

    return ImportResult(chapters=chapters, settings=settings, materials=materials)


def _is_setting(text: str) -> bool:
    """判断是否为设定类内容。"""
    return any(kw in text for kw in _SETTING_KEYWORDS)


def _looks_like_chapter(title: str, text: str) -> bool:
    """判断 section 是否像章节（而非设定/材料）。"""
    # 标题含章节分隔符
    for pattern in _CHAPTER_PATTERNS:
        if pattern.search(title):
            return True

    # 标题是「第X章」形式
    if re.match(r'第[一二三四五六七八九十百千\d]+[章节]', title):
        return True

    # 内容以「第X章」开头
    first_100 = text[:100]
    for pattern in _CHAPTER_PATTERNS:
        if pattern.search(first_100):
            return True

    return False


def _auto_tag(text: str, title: str) -> list[str]:
    """根据内容关键词自动生成标签。"""
    combined = (text[:300] + " " + title).lower()
    tags: list[str] = []
    for tag, keywords in _MATERIAL_TAG_HINTS.items():
        if any(kw.lower() in combined for kw in keywords):
            tags.append(tag)
    return tags if tags else ["未分类"]


# ═══════════════════════════════════════════════════════════════
# Final Assembly
# ═══════════════════════════════════════════════════════════════

def _finalize_import(
    sections_raw: list[_Section],
    fpath: Path,
    file_size: int,
    fmt: str,
) -> ImportResult:
    """组装最终 ImportResult 并填入元数据。"""
    result = _classify_sections(sections_raw)

    full_text = "\n".join(
        s.title + "\n" + s.content for s in sections_raw
    )

    result.metadata.update({
        "source_filename": fpath.name,
        "total_chars": len(full_text),
        "detected_format": fmt,
        "word_count": _count_chars(full_text),
        "file_size": _readable_size_bytes(file_size),
    })

    if file_size > DocumentParser.LARGE_FILE_THRESHOLD:
        logger.warning(
            "DocumentParser: file '%s' is %.1fKB (>100KB)",
            fpath.name, file_size / 1024,
        )

    return result


# ═══════════════════════════════════════════════════════════════
# Generic Helpers
# ═══════════════════════════════════════════════════════════════

def _count_chars(text: str) -> int:
    """统计字符数（去换行）。"""
    return len(text.replace("\n", "").replace("\r", ""))


def _readable_size_bytes(byte_size: int) -> str:
    """字节数 → 可读字符串。"""
    for unit in ["B", "KB", "MB"]:
        if byte_size < 1024:
            return f"{byte_size:.0f}{unit}"
        byte_size /= 1024
    return f"{byte_size:.1f}GB"
